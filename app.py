"""
剧本拆解大师v2.52版 - Flask Web 应用
支持上传/粘贴剧本，通过 AI 进行角色、道具、场景、分镜四步提取
"""
# (C) foxpaw, 2026-07-15


import os
import sys
import json
import time
import uuid
import logging
import re
import subprocess
import threading
import requests
from flask import (
    Flask, request, jsonify, send_file,
    Response, stream_with_context
)
from werkzeug.utils import safe_join

import config
from utils import get_base_path
from extractors import characters, props, scenes, shots, emotion_timeline

# 导入拆分后的模块
from services.ai_service import call_ai
from services.file_parser import parse_script, validate_file_type
from utils.text import detect_episode, split_script_by_episodes
from utils.url import validate_api_base_url
from utils.validation import sanitize_api_config
from asset_audit.auditor import extract_assets, merge_assets
from asset_audit.pptx_builder import build_episode_pptx
from asset_audit.html_builder import build_summary_html
from utils.sse import json_sse
from reports.word_report import generate_word_report
from reports.html_report import generate_html_report


# 道具提示词固定前段（代码拼接，AI不输出）
PROPS_PROMPT_PREFIX = (
    '道具设计图，三视图展示(正面视图、侧面视图、俯视俯拍)+超特写纹理细节，'
    '纯白背景，产品摄影风格，'
)

# 道具提示词固定后段（代码拼接，AI不输出，确保永不遗漏）
PROPS_PROMPT_SUFFIX = (
    'Netflix剧集级道具摄影质感，Arri Alexa摄影机拍摄，HDR高动态范围影像，'
    '精致色彩分级，画面通透干净，亮部细节丰富，暗部层次清晰，柔和对比度，'
    '电影级构图，真实物理材质呈现。 写实超写实质感，真实物理材质与表面纹理，'
    '自然使用痕迹与磨损，真实物理光影，8K超高清，极致细腻真实纹理，'
    '无损高清画质，锐利真实细节。 材质统一一致，各视角相同纹理，'
    '相同配色，各视图比例一致，年代风格统一，磨损程度一致，文字标识一致。 '
    '不要出现：AI生成质感，3D渲染质感，CGI特效感，游戏引擎画面，卡通渲染，'
    '手绘插画，动画风格，皮克斯风格，塑料/蜡质/瓷娃娃质感，光滑完美无瑕疵的CG材质，'
    '无使用痕迹的完美表面，美颜滤镜，磨皮过度，过度锐化，数字平滑感，过度干净渲染，'
    '高光油腻，非自然均匀光照，合成光照，非自然调色，卡通色彩，纹理重复，'
    '漂浮镜头，机械相机运动，完美稳定，人工帧插值，超真实渲染，非自然锐利边缘。'
)

# 人物提示词固定前段（代码拼接，AI不输出）
CHARACTER_PROMPT_PREFIX = (
    '画面左侧是人物正面锁骨以上全脸特写+画面右侧是人物三视图'
    '(正面视角、正侧面视角、背面视角)，全身站立姿态，双手自然下垂，表情自然，'
)

# 人物提示词固定后段（代码拼接，AI不输出，确保永不遗漏）
CHARACTER_PROMPT_SUFFIX = (
    'realistic skin texture with visible pores, subtle natural skin oil, '
    'realistic subsurface scattering, fine facial details, naturally backlit peach fuzz, '
    'authentic skin imperfections, cinematic realism, organic rendering。 '
    'Netflix剧集级角色肖像质感，Arri Alexa摄影机拍摄，HDR高动态范围影像，'
    '精致色彩分级，画面通透干净，亮部细节丰富不死白，暗部层次清晰不死黑，'
    '柔和对比度，均匀自然曝光，4K/8K超高清分辨率，电影级构图，自然肤色还原。 '
    '写实超写实质感，真实物理材质与皮肤纹理，自然瑕疵与使用痕迹，'
    '真实物理光影，极致细腻真实纹理，浅景深虚化背景，立体空间纵深感，'
    '浓厚叙事氛围感，无损高清画质，锐利真实细节。 '
    '同一个人物角色，面部特征统一一致，全程穿着同一套服装，服装细节完全相同，'
    '相同体型，体格特征一致，保持相同年龄面貌，发型一致，发色统一，文字标识一致。 '
    '反向提示词：不要出现AI生成面孔，CGI质感，游戏引擎画面，塑料/蜡质皮肤，'
    '瓷娃娃皮肤，美颜滤镜，磨皮过度，过度锐化，数字平滑感，过度干净渲染，'
    '高光油腻，非自然面部比例，完美对称，合成光照，均匀光照，非自然眼部反光，'
    '虚假景深，人工电影模糊，非自然调色，卡通色彩，纹理重复，漂浮镜头，'
    '机械相机运动，完美稳定，人工帧插值，超真实渲染，非自然锐利边缘。'
)


# API 配置参数键列表（统一引用，避免重复定义）
API_CONFIG_KEYS = ["provider", "api_key", "model", "base_url", "temperature", "top_p", "frequency_penalty", "presence_penalty", "max_tokens", "thinking"]

app = Flask(__name__,
    static_folder=os.path.join(get_base_path(), 'static'),
    static_url_path='',
    template_folder=os.path.join(get_base_path(), 'templates'))
app.secret_key = config.SECRET_KEY
app.config['MAX_CONTENT_LENGTH'] = config.MAX_CONTENT_LENGTH
app.config['UPLOAD_FOLDER'] = config.UPLOAD_FOLDER

# ===== 安全增强: CORS + 安全响应头 =====
@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Cache-Control'] = 'no-store'
    # CORS: 仅允许同源或配置白名单
    origin = request.headers.get('Origin')
    if origin:
        allowed_origins = config.ALLOWED_ORIGINS
        if not allowed_origins:
            allowed_origins = [request.host_url.rstrip('/')]
        if origin in allowed_origins:
            response.headers['Access-Control-Allow-Origin'] = origin
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

# ===== 安全增强: 速率限制 =====
_rate_limit_store = {}

def _check_rate_limit(key, max_requests=30, window=60):
    """简单的内存速率限制: 每窗口最多 max_requests 次"""
    now = time.time()
    if key in _rate_limit_store:
        count, window_start = _rate_limit_store[key]
        if now - window_start > window:
            _rate_limit_store[key] = (1, now)
            return True
        if count >= max_requests:
            return False
        _rate_limit_store[key] = (count + 1, window_start)
        return True
    _rate_limit_store[key] = (1, now)
    return True

def _get_client_id():
    """获取客户端标识用于速率限制"""
    return request.headers.get('X-Forwarded-For', request.remote_addr or 'unknown')

os.makedirs(config.UPLOAD_FOLDER, exist_ok=True)


def _series_name(script_name):
    """从剧本名提取系列名：移除 _第X集 后缀，实现单系列共享知识库"""
    if not script_name:
        return script_name
    import re
    # 匹配 _第X集 或 _第XX集 或 第X集（无下划线时也考虑）
    cleaned = re.sub(r'_第[一二三四五六七八九十\d]+集$', '', script_name)
    # 如果没变化（没有 _第X集），也尝试去掉末尾的 第X集
    if cleaned == script_name:
        cleaned = re.sub(r'第[一二三四五六七八九十\d]+集$', '', script_name)
    cleaned = cleaned.strip(' _-')
    # 如果去掉后缀后为空（如用户直接命名"第1集"），回退到 script_name 自身
    if not cleaned:
        return script_name
    return cleaned


def _load_temp_knowledge(script_name):
    """加载剧本的临时知识库（自动从磁盘读取，优先系列级目录）"""
    import os, json
    from werkzeug.utils import safe_join
    
    series = _series_name(script_name)
    # 先尝试系列级目录
    d = safe_join(config.OUTPUT_DIR, series) if series else safe_join(config.OUTPUT_DIR, script_name) if script_name else config.OUTPUT_DIR
    if not d:
        return None
    os.makedirs(d, exist_ok=True)
    kb_path = os.path.join(d, f'{series}_人物小传大纲_知识库.json')
    if os.path.exists(kb_path):
        try:
            with open(kb_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            return None
    # 兼容旧路径：旧版 kb 在脚本级目录
    d2 = safe_join(config.OUTPUT_DIR, script_name) if script_name else None
    if d2:
        kb_path2 = os.path.join(d2, f'{script_name}_人物小传大纲_知识库.json')
        if os.path.exists(kb_path2):
            try:
                with open(kb_path2, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except (FileNotFoundError, json.JSONDecodeError):
                return None
    return None
def _output_path(script_name, filename, episode_label=''):
    """获取剧本专属输出路径。episode_label 非空时在系列目录下创建剧集子目录。"""
    if not script_name:
        d = config.OUTPUT_DIR
    else:
        series = _series_name(script_name)
        d = safe_join(config.OUTPUT_DIR, series) if series else safe_join(config.OUTPUT_DIR, script_name)
    if not d:
        raise ValueError(f"非法剧本名称: {script_name}")
    if episode_label:
        d = os.path.join(d, episode_label)
    os.makedirs(d, exist_ok=True)
    return os.path.join(d, filename)

def _map_extraction_to_pptx_assets(extraction_results):
    """将分步提取结果映射为 PPTX 生成器所需的字段格式"""
    mapped_chars = []
    for c in extraction_results.get('characters', {}).get('characters', []):
        desc = (c.get('description', '') or '').strip()
        if len(desc) > 20:
            desc = desc[:18] + '..'
        mapped_chars.append({
            'name_cn': c.get('name', ''),
            'name_en': '',
            'costume': desc,
        })
    mapped_props = []
    for p in extraction_results.get('props', {}).get('props', []):
        mapped_props.append({
            'name_cn': p.get('name', ''),
            'usage': p.get('usage', ''),
        })
    mapped_scenes = []
    for s in extraction_results.get('scenes', {}).get('scenes', []):
        mapped_scenes.append({
            'name_cn': s.get('title', ''),
            'name_en': s.get('location', ''),
            'synopsis': s.get('synopsis', ''),
        })
    return {
        'characters': mapped_chars,
        'props': mapped_props,
        'scenes': mapped_scenes,
    }


def _cleanup_intermediate_files(script_name, episode_label):
    """Delete intermediate step HTML files after final report is generated."""
    intermediate_files = ['角色提取.html', '道具提取.html', '场景拆解.html', '分镜拆解.html']
    series = _series_name(script_name)
    d = safe_join(config.OUTPUT_DIR, series, episode_label) if series else safe_join(config.OUTPUT_DIR, script_name, episode_label)
    if not d:
        return
    for fname in intermediate_files:
        fpath = os.path.join(d, fname)
        if os.path.isfile(fpath):
            try:
                os.remove(fpath)
            except OSError:
                pass



# ===== 全剧进度文件（analysis_progress.json）=====

import datetime as _dt

def _progress_path(script_name):
    """Get analysis_progress.json path for a series"""
    return _output_path(script_name, 'analysis_progress.json')


def _load_progress(script_name):
    """Load full analysis progress, or None"""
    try:
        pp = _progress_path(script_name)
        if os.path.isfile(pp):
            with open(pp, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        logging.warning("进度文件读取失败: %s", str(e), exc_info=True)
    return None


def _save_progress(script_name, progress_data):
    """Save full analysis progress"""
    try:
        pp = _progress_path(script_name)
        os.makedirs(os.path.dirname(pp), exist_ok=True)
        progress_data['updated_at'] = _dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        with open(pp, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.warning("进度文件保存失败: %s", str(e), exc_info=True)


def _init_progress(script_name, episodes, selected_eps=None):
    """Create initial progress file after split"""
    now = _dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    ep_map = {}
    for ep in episodes:
        ep_map[ep['label']] = {
            'status': 'pending',
            'episode_num': ep.get('episode', 0),
            'steps': {
                'characters': {'status': 'pending'},
                'props': {'status': 'pending'},
                'scenes': {'status': 'pending'},
                'shots': {'status': 'pending'},
            }
        }
    data = {
        'script_name': script_name,
        'total_episodes': len(episodes),
        'selected_episodes': selected_eps or [ep.get('episode', 0) for ep in episodes],
        'status': 'running',
        'started_at': now,
        'updated_at': now,
        'episodes': ep_map,
    }
    _save_progress(script_name, data)


def _update_episode_step(script_name, episode_label, step_name, step_status, **kwargs):
    """Update a single step status in progress file"""
    prog = _load_progress(script_name)
    if not prog:
        return
    eps = prog.get('episodes', {})
    ep = eps.get(episode_label)
    if not ep:
        return
    ep['status'] = 'running'
    now = _dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if not ep.get('started_at'):
        ep['started_at'] = now
    step = ep['steps'].get(step_name)
    if step:
        step['status'] = step_status
        if step_status == 'processing' and not step.get('started_at'):
            step['started_at'] = now
        if step_status in ('completed', 'failed'):
            step['completed_at'] = now
        for k, v in kwargs.items():
            step[k] = v
    _save_progress(script_name, prog)


def _mark_episode_complete(script_name, episode_label):
    """Mark an episode as completed in progress"""
    prog = _load_progress(script_name)
    if not prog:
        return
    eps = prog.get('episodes', {})
    ep = eps.get(episode_label)
    if ep:
        ep['status'] = 'completed'
        ep['completed_at'] = _dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    # Check if all selected episodes are done
    selected = [str(e) for e in prog.get('selected_episodes', [])]
    all_done = all(
        eps.get(lbl, {}).get('status') == 'completed'
        for lbl in eps
        if lbl.startswith('EPISODE') and any(lbl.endswith(f' {e}') or lbl.endswith(f' {str(e).zfill(2)}') for e in selected)
    )
    # Simpler check: all episodes with label in progress are completed
    all_eps_done = all(eps[lbl].get('status') == 'completed' for lbl in eps if eps[lbl].get('episode_num', 0) in prog.get('selected_episodes', []))
    if all_eps_done:
        prog['status'] = 'completed'
    _save_progress(script_name, prog)


def _is_episode_completed(script_name, episode_label):
    """Check if an episode is fully completed"""
    prog = _load_progress(script_name)
    if not prog:
        return False
    ep = prog.get('episodes', {}).get(episode_label)
    return ep and ep.get('status') == 'completed'


def _get_completed_steps(script_name, episode_label):
    """Get list of completed step names for an episode"""
    prog = _load_progress(script_name)
    if not prog:
        return []
    ep = prog.get('episodes', {}).get(episode_label)
    if not ep:
        return []
    return [s for s, d in ep.get('steps', {}).items() if d.get('status') == 'completed']



def _call_ai_retry(sys_p, user_p, api_config, step_label, max_retries=2, temp=None):
    """Call AI with retry. Returns (result, None) on success, (None, error_msg) on final failure."""
    cfg = {**api_config} if temp is not None else api_config
    if temp is not None:
        cfg['temperature'] = str(temp)
    last_error = None
    for attempt in range(max_retries):
        try:
            result = call_ai(sys_p, user_p, cfg)
            return result, None
        except Exception as e:
            last_error = str(e)
            if attempt < max_retries - 1:
                time.sleep(2 ** min(attempt, 3))
    return None, f"{step_label}已重试{max_retries}次仍失败: {last_error}"


# ============================================================
# 路由
# ============================================================


@app.route('/')
def index():
    return app.send_static_file('index.html')


@app.route('/api/config')
def get_config():
    """获取支持的 AI 提供商配置"""
    providers = {}
    for key, info in config.AI_PROVIDERS.items():
        providers[key] = {
            "name": info["name"],
            "default_model": info["default_model"],
        }
    return jsonify({"providers": providers})


@app.route('/api/list_models', methods=['POST'])
def list_models():
    """调用 DeepSeek API 获取可用模型列表"""
    if not _check_rate_limit(_get_client_id(), max_requests=30, window=60):
        return jsonify({"error": "请求过于频繁，请稍后重试"}), 429
    data = request.get_json(silent=True) or {}
    api_key = data.get('api_key', '')
    base_url = data.get('base_url', config.DEEPSEEK_API_URL)
    try:
        validate_api_base_url(base_url, 'deepseek')
    except ValueError as e:
        return jsonify({
            "error": str(e),
            "models": config.DEEPSEEK_KNOWN_MODELS,
            "source": "fallback"
        }), 200

    if not api_key:
        # 无 API Key 时返回后备模型列表
        return jsonify({
            "models": config.DEEPSEEK_KNOWN_MODELS,
            "source": "fallback"
        })
    
    try:
        headers = {
            "Accept": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        resp = requests.get(
            f"{base_url.rstrip('/')}/models",
            headers=headers,
            timeout=10
        )
        resp.raise_for_status()
        data = resp.json()
        
        models = data.get('data', [])
        # 过滤和格式化
        formatted = []
        for m in models:
            model_id = m.get('id', '')
            # 排除 embedding 等非对话模型
            if 'embed' in model_id.lower() or 'image' in model_id.lower():
                continue
            formatted.append({
                "id": model_id,
                "name": m.get('id', model_id),
                "owned_by": m.get('owned_by', 'deepseek'),
                "ctx": "1M" if 'v4' in model_id else ("64K" if 'chat' in model_id or 'reasoner' in model_id else "?"),
            })
        
        return jsonify({"models": formatted, "source": "api"})
        
    except requests.exceptions.HTTPError as e:
        status = e.response.status_code
        if status == 401:
            return jsonify({
                "error": "API Key 无效，请检查",
                "models": config.DEEPSEEK_KNOWN_MODELS,
                "source": "fallback"
            }), 200
        return jsonify({
            "error": f"API 错误 ({status})",
            "models": config.DEEPSEEK_KNOWN_MODELS,
            "source": "fallback"
        }), 200
    except requests.exceptions.ConnectionError:
        return jsonify({
            "error": f"无法连接到 {base_url}",
            "models": config.DEEPSEEK_KNOWN_MODELS,
            "source": "fallback"
        }), 200
    except Exception as e:
        return jsonify({
            "error": str(e),
            "models": config.DEEPSEEK_KNOWN_MODELS,
            "source": "fallback"
        }), 200


@app.route('/api/check_connection', methods=['POST'])
def check_connection():
    """测试 AI API 连通性（代理请求，避免浏览器 CORS 问题）"""
    data = request.get_json(silent=True) or {}
    provider = data.get('provider', 'deepseek')
    api_key = data.get('api_key', '')
    base_url = data.get('base_url', '')

    resolved_url = (base_url or (config.OLLAMA_BASE_URL if provider == 'ollama' else config.DEEPSEEK_API_URL)).rstrip('/')
    try:
        validate_api_base_url(resolved_url, provider)
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)})

    if provider == 'ollama':
        url = resolved_url + '/api/tags'
        try:
            resp = requests.get(url, timeout=5)
            if resp.ok:
                return jsonify({"success": True, "message": "Ollama 连接成功"})
            return jsonify({"success": False, "message": f"Ollama 返回 {resp.status_code}"})
        except requests.exceptions.ConnectionError:
            return jsonify({"success": False, "message": f"无法连接到 Ollama ({url})"})
        except Exception as e:
            return jsonify({"success": False, "message": str(e)})
    else:
        url = resolved_url + '/models'
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        try:
            resp = requests.get(url, headers=headers, timeout=5)
            if resp.ok:
                data = resp.json()
                count = len(data.get('data', []))
                return jsonify({"success": True, "message": f"连接成功! 找到 {count} 个模型", "model_count": count})
            elif resp.status_code == 401:
                return jsonify({"success": False, "message": "API Key 无效"})
            return jsonify({"success": False, "message": f"API 返回 {resp.status_code}"})
        except requests.exceptions.ConnectionError:
            return jsonify({"success": False, "message": f"无法连接到 {url}"})
        except Exception as e:
            return jsonify({"success": False, "message": str(e)})


def _require_auth():
    """可选的 API 认证 — 通过环境变量 AUTH_TOKEN 启用"""
    token = os.environ.get('AUTH_TOKEN', '')
    if not token:
        return True  # 未配置则不启用认证
    auth_header = request.headers.get('Authorization', '')
    if auth_header == f'Bearer {token}':
        return True
    return False


@app.route('/api/analyze', methods=['POST'])
def analyze():
    """主分析接口：接收剧本并运行四步提取（SSE 流式进度）"""
    if not _require_auth():
        return jsonify({"error": "未授权访问"}), 401
    if not _check_rate_limit(_get_client_id(), max_requests=20, window=60):
        return jsonify({"error": "请求过于频繁，请稍后重试"}), 429
    
    def generate():
        results = {}
        error = None
        
        try:
            # 1. 获取剧本内容
            script_text = ""
            script_name = "未命名剧本"
            
            if 'file' in request.files:
                file = request.files['file']
                if file.filename == '':
                    yield json_sse("error", {"message": "请选择要上传的文件"})
                    return
                
                # MIME 类型校验
                ext = os.path.splitext(file.filename)[1].lower()
                ALLOWED_EXTS = {'.pdf': 'application/pdf', '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.doc': 'application/msword', '.txt': 'text/plain', '.md': 'text/markdown', '.markdown': 'text/markdown'}
                if ext not in ALLOWED_EXTS:
                    yield json_sse("error", {"message": f"不支持的文件格式: {ext}，请上传 PDF、Word、TXT 或 MD 文件"})
                    return
                
                # 保存上传文件
                safe_name = f"{uuid.uuid4().hex}{ext}"
                file_path = os.path.join(config.UPLOAD_FOLDER, safe_name)
                file.save(file_path)

                # 校验文件内容是否与扩展名一致
                if not validate_file_type(file_path):
                    try:
                        os.remove(file_path)
                    except OSError:
                        pass
                    yield json_sse("error", {"message": "文件内容与实际格式不符，请检查文件"})
                    return

                try:
                    script_text = parse_script(file_path)
                except Exception as e:
                    logging.getLogger("app").error("文件解析失败: %s", str(e))
                    yield json_sse("error", {"message": f"文件解析失败: {str(e)}"})
                    return
                finally:
                    # 清理临时文件
                    try:
                        os.remove(file_path)
                    except OSError:
                        pass
                
                script_name = os.path.splitext(file.filename)[0]
            else:
                data = request.get_json(silent=True)
                if data and 'text' in data:
                    script_text = data['text']
                elif request.form.get('text'):
                    script_text = request.form['text']
                if data and 'script_name' in data:
                    script_name = data['script_name']
                elif request.form.get('script_name'):
                    script_name = request.form['script_name']
            
            # 净化文件名（防止路径穿越）
            script_name = re.sub(r'[\\/:*?"<>|]', '_', script_name)
            script_name = script_name.replace('\x00', '')
            script_name = script_name.lstrip('.-')
            script_name = script_name.strip()[:100] or "未命名剧本"
            
            if not script_text or not script_text.strip():
                yield json_sse("error", {"message": "剧本内容为空，请提供有效的剧本"})
                return
            
            # 截断太长文本（50000字足够覆盖大部分剧本单幕或短剧全本）
            if len(script_text) > 50000:
                yield json_sse("info", {"message": f"剧本较长({len(script_text)}字符)，将完整处理而非截断"})
            # 不再截断剧本内容，完整传递给AI
            
            # 获取 AI 配置（从 FormData 中提取）
            api_config = {}
            for key in ['provider', 'api_key', 'model', 'base_url', 'temperature', 'top_p', 'frequency_penalty', 'presence_penalty', 'max_tokens', 'thinking']:
                val = request.form.get(f'api_config[{key}]')
                if val is not None:
                    api_config[key] = val
            api_config = sanitize_api_config(api_config)

            # 获取分析步骤参数（空=全部，指定则只跑该步）
            step_filter = request.form.get('step', '').strip()
            resume_step = request.form.get('resume_step', '').strip()
            
            # 批量模式参数
            batch_mode = request.form.get('batch_mode', '').strip()
            episode_label = request.form.get('episode_label', '').strip()
            # 批量模式下输出子目录（如 EPISODE 01）
            ep_subdir = episode_label if batch_mode == '1' and episode_label else ''
            
            # 断点续传：从 progress 文件加载已完成的步骤
            completed_steps = _get_completed_steps(script_name, ep_subdir) if ep_subdir else []
            if completed_steps:
                yield json_sse("info", {"message": f"检测到已完成步骤：{', '.join(completed_steps)}，从断点继续..."})
                pass  # resume_step is used below
            if completed_steps:
                yield json_sse("info", {"message": f"检测到已完成步骤：{', '.join(completed_steps)}，从断点继续..."})
            
            # 读取剧本拆解风格
            breakdown_style = request.form.get('breakdown_style', 'normal').strip()
            if breakdown_style not in ('female', 'male', 'normal'):
                breakdown_style = 'normal'
            
            # 加载或自动生成临时知识库
            temp_kb = _load_temp_knowledge(script_name)
            
            # 自动预处理：知识库不存在但有材料时，自动生成
            char_bio_in = request.form.get('char_bio', '').strip()
            outline_in = request.form.get('story_outline', '').strip()
            needs_materials = (char_bio_in or outline_in) and not temp_kb
            if needs_materials:
                pre_sys = """你是剧本分析预处理专家。你的任务是分析人物小传和故事大纲，提取结构化信息。

人物小传分析要求：
- 提取每个角色的名称(name)、别名(aliases)、年龄种族(age_race)、性格标签(personality)、弧光走向(arc)、外貌特征(appearance)、人际关系(relationships)、关键特征(key_traits)

故事大纲分析要求：
- 提取年代/时空设定(time_period)、主要地点(location)、类型标签(genre)、整体基调(tone)、关键剧情道具(key_props)、关键情节点(key_plot_points)、环境风格描述(environment_style)

输出纯JSON："""
                pre_user = "请分析以下内容并生成结构化知识库。\n\n"
                if char_bio_in:
                    pre_user += f"=== 人物小传 ===\n{char_bio_in}\n\n"
                if outline_in:
                    pre_user += f"=== 故事大纲 ===\n{outline_in}\n\n"
                pre_user += (
                    '输出JSON格式：\n'
                    '{\n'
                    '  "characters": [{\n'
                    '    "name": "角色名",\n'
                    '    "aliases": ["别名"],\n'
                    '    "age_race": "年龄/种族",\n'
                    '    "personality": ["性格标签"],\n'
                    '    "arc": "角色弧光走向",\n'
                    '    "appearance": "外貌特征",\n'
                    '    "relationships": [{"target": "关联角色", "relation": "关系描述"}],\n'
                    '    "key_traits": "关键特征"\n'
                    '  }],\n'
                    '  "world": {\n'
                    '    "time_period": "年代/时空设定",\n'
                    '    "location": "主要地点",\n'
                    '    "genre": "类型标签",\n'
                    '    "tone": "整体基调",\n'
                    '    "key_props": ["关键剧情道具"],\n'
                    '    "key_plot_points": ["关键情节点"],\n'
                    '    "environment_style": "环境风格描述"\n'
                    '  }\n'
                    '}\n'
                    '只输出JSON。'
                )
                try:
                    api_cfg = {}
                    for key in ['provider', 'api_key', 'model', 'base_url', 'temperature', 'top_p', 'frequency_penalty', 'presence_penalty', 'max_tokens', 'thinking']:
                        val = request.form.get(f'api_config[{key}]', '')
                        if val:
                            api_cfg[key] = val
                    api_cfg = sanitize_api_config(api_cfg)
                    raw = call_ai(pre_sys, pre_user, api_cfg)
                    json_text = raw.strip()
                    js = json_text.find('{')
                    je = json_text.rfind('}')
                    if js != -1 and je != -1:
                        json_text = json_text[js:je+1]
                    parsed = json.loads(json_text)
                    series = _series_name(script_name)
                    d = safe_join(config.OUTPUT_DIR, series) if series else safe_join(config.OUTPUT_DIR, script_name)
                    if d:
                        os.makedirs(d, exist_ok=True)
                        kb_path = os.path.join(d, f'{series}_人物小传大纲_知识库.json')
                        with open(kb_path, 'w', encoding='utf-8') as f:
                            json.dump(parsed, f, ensure_ascii=False, indent=2)
                        temp_kb = parsed
                        yield json_sse("info", {"message": f"已自动分析人物小传和故事大纲（{len(parsed.get('characters', []))}个角色）"})
                except Exception as e:
                    logging.getLogger("app").warning("自动预处理失败: %s", str(e))
                    yield json_sse("info", {"message": "人物小传分析失败，将以普通模式继续"})
            
            # 检查是否需要提示用户（无材料且无kb）
            has_any_material = bool(char_bio_in or outline_in)
            no_material_flag = request.form.get('no_materials_skip', '0')
            
            # 加载已有结果（分步模式用）
            prior_json = request.form.get('prior_results', '')
            if prior_json:
                try:
                    prior = json.loads(prior_json)
                    for k in ['characters', 'props', 'scenes', 'shots']:
                        if k in prior:
                            results[k] = prior[k]
                except Exception:
                    logging.warning("prior_results 解析失败，跳过", exc_info=True)
            
            # 预分块：长剧本切成小块用于第一轮扫描
            CHUNK_SIZE = 4000  # 约1000个中文字符，平衡分块粒度与AI上下文质量
            script_chunks = []
            if len(script_text) > CHUNK_SIZE:
                paragraphs = script_text.split('\n\n')
                current = ""
                for p in paragraphs:
                    if len(current) + len(p) > CHUNK_SIZE and current:
                        script_chunks.append(current.strip())
                        current = p
                    else:
                        current += ("\n\n" if current else "") + p
                if current.strip():
                    script_chunks.append(current.strip())
            else:
                script_chunks = [script_text]
            
            # 2. 分析接口调用（可选的）
            yield json_sse("info", {
                "message": f"开始分析《{script_name}》，剧本长度：{len(script_text)} 字符",
                "script_length": len(script_text)
            })
            
            # ===== 步骤1: 角色提取 =====
            if (not step_filter or step_filter == 'characters') and 'characters' not in completed_steps:
                if ep_subdir:
                    _update_episode_step(script_name, ep_subdir, 'characters', 'processing')
                yield json_sse("progress", {"step": "characters", "status": "processing", "label": "角色提取", "message": "第一轮：全剧本扫描角色名..."})
                try:
                    # 全剧本一次发送，利用 DeepSeek V4 1M 上下文，避免分块撕裂角色信息
                    sys_p, user_p = characters.build_list_prompt(script_text)
                    raw, retry_err = _call_ai_retry(sys_p, user_p, api_config, '角色提取', max_retries=3, temp=0)
                    if retry_err:
                        yield json_sse('pause', {'step': 'characters', 'message': retry_err, 'data': {'results': results}})
                        return
                    names = characters.parse_list(raw)
                    all_names = set(names)
                    
                    char_names = sorted(all_names)
                    if not char_names:
                        raise RuntimeError("未找到任何角色名")
                    
                    yield json_sse("progress", {"step": "characters", "status": "processing", "label": "角色提取", "message": f"发现 {len(char_names)} 个角色，第二轮：生成详情和提示词..."})
                    sys_p, user_p = characters.build_detail_prompt(script_text, char_names, temp_kb=temp_kb)
                    raw, retry_err = _call_ai_retry(sys_p, user_p, api_config, '角色提取第二轮')
                    if retry_err:
                        yield json_sse('pause', {'step': 'characters', 'message': retry_err, 'data': {'results': results}})
                        return
                    result = characters.parse_result(raw)
                    for c in result.get('characters', []):
                        if c.get('prompt'):
                            c['prompt'] = CHARACTER_PROMPT_PREFIX + c['prompt'] + CHARACTER_PROMPT_SUFFIX
                    results['characters'] = result
                    cc = len(result.get('characters', []))
                    if cc == 0:
                        raw_preview = str(raw)[:400] if raw else '(AI未返回内容)'
                        logging.getLogger("app").warning("角色第二轮返回0个角色，AI原始返回前500字: %s", str(raw)[:500])
                        raise RuntimeError(f"角色详情生成失败：第二轮返回0个角色。AI原始返回前400字: {raw_preview}")
                    # 按角色重要性排序：主角 > 配角 > 龙套 > 未知
                    char_order = {'主角': 0, '配角': 1, '龙套': 2}
                    result['characters'] = sorted(result.get('characters', []),
                        key=lambda c: char_order.get(c.get('role_type', ''), 99))
                    # 生成分步 HTML 报告
                    partial_html = generate_html_report(results, script_name, episode_info=None)
                    partial_path = _output_path(script_name, '角色提取.html', ep_subdir)
                    with open(partial_path, 'w', encoding='utf-8') as f:
                        f.write(partial_html)
                    if ep_subdir:
                        _update_episode_step(script_name, ep_subdir, 'characters', 'completed', result_summary=f'识别 {cc} 个角色')
                    yield json_sse("progress", {"step": "characters", "status": "complete", "label": "角色提取", "message": f"完成！识别 {cc} 个角色", "data": result, "download_url": f"/api/download/{os.path.basename(partial_path)}"})
                except Exception as e:
                    yield json_sse("progress", {"step": "characters", "status": "error", "label": "角色提取", "message": f"失败：{str(e)}"})
                    if step_filter and step_filter == 'characters':
                        yield json_sse("error", {"message": f"角色提取失败: {str(e)}"})
                        return
                    results['characters'] = {"characters": [], "total_count": 0, "error": str(e)}
            
            # ===== 步骤2: 道具提取 =====
            if (not step_filter or step_filter == 'props') and 'props' not in completed_steps:
                if ep_subdir:
                    _update_episode_step(script_name, ep_subdir, 'props', 'processing')
                # 检测集数：单集无频率限制，多集要求≥2场
                eps = split_script_by_episodes(script_text)
                ep_count = len(eps) if eps else 0
                min_freq = 1
                freq_hint = "单集模式：提取全部道具" if min_freq <= 1 else f"多集模式（{ep_count}集）：仅提取≥2场道具"
                yield json_sse("progress", {"step": "props", "status": "processing", "label": "道具提取", "message": f"第一轮：扫描全剧本道具名（{freq_hint}）..."})
                try:
                    prop_names = []
                    for attempt in range(2):
                        sys_p, user_p = props.build_list_prompt(script_text, min_appearances=min_freq)
                        if attempt > 0:
                            user_p = "（重试，请务必列出所有道具名称）\n" + user_p
                        raw, retry_err = _call_ai_retry(sys_p, user_p, api_config, '道具提取' + ('(重试)' if attempt > 0 else ''), max_retries=3, temp=0)
                        if retry_err:
                            # already retried at _call_ai_retry level, give up
                            pass
                        prop_names = props.parse_list(raw) if raw else []
                        if not isinstance(prop_names, list):
                            prop_names = []
                        if len(prop_names) > 0:
                            break
                        if attempt == 0:
                            yield json_sse("progress", {"step": "props", "status": "processing", "label": "道具提取", "message": "第一轮返回空，重试中..."})
                            time.sleep(2)
                    
                    if len(prop_names) == 0:
                        raw_preview = str(raw)[:300] if raw else '(AI未返回内容)'
                        logging.getLogger("app").warning("道具第一轮返回0个道具，AI原始返回前500字: %s", str(raw)[:500])
                        raise RuntimeError(f"道具第一轮未识别到任何道具（已重试1次）。AI原始返回: {raw_preview}")
                    else:
                        yield json_sse("progress", {"step": "props", "status": "processing", "label": "道具提取", "message": f"发现 {len(prop_names)} 个道具，第二轮：生成详情和提示词..."})
                        sys_p, user_p = props.build_detail_prompt(script_text, prop_names, results, temp_kb=temp_kb)
                        raw, pretry_err = _call_ai_retry(sys_p, user_p, api_config, '道具提取第二轮')
                        if pretry_err:
                            yield json_sse('pause', {'step': 'props', 'message': pretry_err, 'data': {'results': results}})
                            return
                        result = props.parse_result(raw)
                        # 代码拼接固定后段，AI只输出变量部分
                        for p in result.get('props', []):
                            if p.get('prompt'):
                                p['prompt'] = PROPS_PROMPT_PREFIX + p['prompt'] + PROPS_PROMPT_SUFFIX
                        results['props'] = result
                        pc = len(result.get('props', []))
                        if pc == 0:
                            yield json_sse("progress", {"step": "props", "status": "complete", "label": "道具提取", "message": f"⚠️ 第二轮解析失败。AI原始返回前500字: {raw[:500]}"})
                        else:
                            # 生成分步 HTML 报告
                            partial_html = generate_html_report(results, script_name, episode_info=None)
                            partial_path = _output_path(script_name, '道具提取.html', ep_subdir)
                            with open(partial_path, 'w', encoding='utf-8') as f:
                                f.write(partial_html)
                            if ep_subdir:
                                _update_episode_step(script_name, ep_subdir, 'props', 'completed', result_summary=f'识别 {pc} 个道具')
                            yield json_sse("progress", {"step": "props", "status": "complete", "label": "道具提取", "message": f"完成！识别 {pc} 个道具", "data": result, "download_url": f"/api/download/{os.path.basename(partial_path)}"})
                except Exception as e:
                    yield json_sse("progress", {"step": "props", "status": "error", "label": "道具提取", "message": f"失败：{str(e)}"})
                    if step_filter and step_filter == 'props':
                        yield json_sse("error", {"message": f"道具提取失败: {str(e)}"})
                        return
                    results['props'] = {"props": [], "total_count": 0, "error": str(e)}
            
            # ===== 步骤3: 场景拆解 =====
            if (not step_filter or step_filter == 'scenes') and 'scenes' not in completed_steps:
                if ep_subdir:
                    _update_episode_step(script_name, ep_subdir, 'scenes', 'processing')
                yield json_sse("progress", {"step": "scenes", "status": "processing", "label": "场景拆解", "message": "第一轮：识别所有场景场次..."})
                try:
                    sys_p, user_p = scenes.build_list_prompt(script_text)
                    raw, sretry_err = _call_ai_retry(sys_p, user_p, api_config, '场景拆解', max_retries=3, temp=0)
                    if sretry_err:
                        yield json_sse('pause', {'step': 'scenes', 'message': sretry_err, 'data': {'results': results}})
                        return
                    scene_list = scenes.parse_list(raw)
                    if not scene_list and raw:
                        logging.getLogger("app").warning("场景第一轮返回0个场景，AI原始返回前500字: %s", str(raw)[:500])
                    if not scene_list:
                        results['scenes'] = {"scenes": [], "total": 0, "summary": "无场景"}
                        yield json_sse("progress", {"step": "scenes", "status": "complete", "label": "场景拆解", "message": "无场景可拆解，跳过"})
                        scene_skip = True
                    else:
                        scene_skip = False
                    
                    if not scene_skip:
                        yield json_sse("progress", {"step": "scenes", "status": "processing", "label": "场景拆解", "message": f"发现 {len(scene_list)} 个场景，第二轮：生成十层描述和双版提示词..."})
                        sys_p, user_p = scenes.build_detail_prompt(script_text, scene_list, results, temp_kb=temp_kb)
                        raw, s2retry_err = _call_ai_retry(sys_p, user_p, api_config, '场景拆解第二轮')
                        if s2retry_err:
                            yield json_sse('pause', {'step': 'scenes', 'message': s2retry_err, 'data': {'results': results}})
                            return
                        result = scenes.parse_result(raw)
                        results['scenes'] = result
                        sc = len(result.get('scenes', []))
                        if sc == 0:
                            raw_preview = str(raw)[:400] if raw else '(AI未返回内容)'
                            logging.getLogger("app").warning("场景第二轮返回0个场景，AI原始返回前500字: %s", str(raw)[:500])
                            raise RuntimeError(f"第二轮生成0个场景详情。AI原始返回前400字: {raw_preview}")
                        # 第三轮：生成多面板布局参考图提示词
                        yield json_sse("progress", {"step": "scenes", "status": "processing", "label": "场景拆解", "message": f"第三轮：生成 {sc} 个场景的多面板布局参考图提示词..."})
                        try:
                            sys_p, user_p = scenes.build_multipanel_prompt(script_text, result, results, temp_kb=temp_kb)
                            raw4, _ = _call_ai_retry(sys_p, user_p, api_config, '场景拆解第三轮')
                            result = scenes.parse_multipanel_result(raw4, result)
                            results['scenes'] = result
                            yield json_sse("progress", {"step": "scenes", "status": "processing", "label": "场景拆解", "message": f"第三轮完成"})
                        except Exception as e4:
                            yield json_sse("progress", {"step": "scenes", "status": "processing", "label": "场景拆解", "message": f"第三轮失败（多面板版将使用自动生成）: {str(e4)[:100]}"})
                            # 兜底：用自动生成填充缺失的 multi_panel_gpt
                            from extractors.scenes import _generate_multi_panel, _fix_lighting_in_mp
                            for s in result.get('scenes', []):
                                if not s.get('multi_panel_gpt'):
                                    s['multi_panel_gpt'] = _generate_multi_panel(s)
                                    s['multi_panel_gpt'] = _fix_lighting_in_mp(s.get('multi_panel_gpt', ''))
                        
                        # 生成分步 HTML 报告
                        partial_html = generate_html_report(results, script_name, episode_info=None)
                        partial_path = _output_path(script_name, '场景拆解.html', ep_subdir)
                        with open(partial_path, 'w', encoding='utf-8') as f:
                            f.write(partial_html)
                        yield json_sse("progress", {"step": "scenes", "status": "complete", "label": "场景拆解", "message": f"完成！拆解 {sc} 个场景", "data": result, "download_url": f"/api/download/{os.path.basename(partial_path)}"})
                except Exception as e:
                    yield json_sse("progress", {"step": "scenes", "status": "error", "label": "场景拆解", "message": f"失败：{str(e)}"})
                    if step_filter and step_filter == 'scenes':
                        yield json_sse("error", {"message": f"场景拆解失败: {str(e)}"})
                        return
                    results['scenes'] = {"scenes": [], "error": str(e)}
            
            # ===== 步骤4: 分镜拆解 =====
            if (not step_filter or step_filter == 'shots') and 'shots' not in completed_steps:
                if ep_subdir:
                    _update_episode_step(script_name, ep_subdir, 'shots', 'processing')
                # 尝试按集拆分剧本
                episodes = split_script_by_episodes(script_text)
                # 如果集拆分失败但剧本很长（>2万字），强制按字数切块
                if (not episodes or len(episodes) <= 2) and len(script_text) > 20000:
                    CHUNK_SIZE = 15000
                    episodes = []
                    pos = 0
                    ep_num = 0
                    while pos < len(script_text):
                        end = min(pos + CHUNK_SIZE, len(script_text))
                        # 尽量在段落边界切断
                        if end < len(script_text):
                            nl = script_text.rfind('\n\n', pos, end)
                            if nl > pos + CHUNK_SIZE // 2:
                                end = nl
                        ep_num += 1
                        episodes.append((ep_num, script_text[pos:end].strip()))
                        pos = end
                    yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"⚠️ 未检测到集标记，已按字数强制分为 {ep_num} 段批次处理"})

                use_batches = episodes and len(episodes) > 2

                if use_batches:
                    batch_size = 1
                    total_episodes = len(episodes)
                    total_batches = (total_episodes + batch_size - 1) // batch_size
                    ep_nums = [str(e[0]) for e in episodes if e[0] > 0]
                    yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"全剧 {total_episodes} 集（检测到：{', '.join(ep_nums[:10])}{'...' if len(ep_nums)>10 else ''}），逐集处理，共 {total_batches} 集..."})

                    all_shot_scenes = []
                    total_shots_all = 0
                    incomplete_batches = []
                    empty_batches = []
                    completed_batches = 0
                    prev_tail = ""

                    for batch_idx in range(0, total_episodes, batch_size):
                        batch_eps = episodes[batch_idx:batch_idx + batch_size]
                        batch_num = batch_idx // batch_size + 1
                        ep_range = f"第{batch_eps[0][0]}集"
                        batch_text = "\n\n".join(t for _, t in batch_eps)
                        batch_result = None

                        # 尝试（最多2次：正常 + 1次重试）
                        for attempt in range(2):
                            retry_tag = "（重试）" if attempt == 1 else ""
                            yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 第一轮规划{retry_tag}..."})
                            try:
                                # 自动启用推理能力（仅模型支持时开启）
                                shots_api_config = dict(api_config)
                                model_name = shots_api_config.get('model', '')
                                if 'v4-flash' in model_name:
                                    shots_api_config['thinking'] = '1'
                                sys_p, user_p = shots.build_list_prompt(batch_text, results, style=breakdown_style)
                                raw, _ = _call_ai_retry(sys_p, user_p, shots_api_config, '分镜拆解规划', max_retries=2)
                                shot_plan = shots.parse_list(raw)

                                yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 第二轮详情{retry_tag}..."})
                                # 情绪时间线预分析
                                et_data = None
                                try:
                                    char_names = [c.get('name','') for c in results.get('characters',{}).get('characters',[])]
                                    scene_list = results.get('scenes',{}).get('scenes',[])
                                    if char_names and scene_list:
                                        yield json_sse("progress", {"step": "shots", "status": "processing", "label": "情绪预分析", "message": "分析角色情绪时间线..."})
                                        et_sys, et_user = emotion_timeline.build_prompt(batch_text, char_names, scene_list)
                                        et_raw, _ = _call_ai_retry(et_sys, et_user, api_config, '情绪时间线', max_retries=2)
                                        et_data = emotion_timeline.parse_result(et_raw)
                                except Exception as e:
                                    logging.warning("情绪时间线分析失败: %s", str(e), exc_info=True)
                                sys_p, user_p = shots.build_detail_prompt(batch_text, shot_plan, results, style=breakdown_style, emotion_timeline=et_data)
                                if prev_tail:
                                    user_p += f"\n\n【上下文衔接】上一批剧情结束于：{prev_tail}。请确保本批首批分镜从该状态自然接续。"
                                raw, _ = _call_ai_retry(sys_p, user_p, api_config, '分镜拆解详情', max_retries=2)
                                batch_result = shots.parse_result(raw)
                                batch_scenes = batch_result.get('scenes', [])
                                batch_shots = batch_result.get('total_shots', 0)
                                if batch_shots == 0:
                                    if attempt == 0:
                                        yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 空结果，重试中...（{raw[:300]}）"})
                                        time.sleep(1)
                                        continue  # 重试
                                    else:
                                        empty_batches.append(ep_range)
                                        yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 重试仍为空 ⚠️（{raw[:300]}）"})
                                else:
                                    all_shot_scenes.extend(batch_scenes)
                                    total_shots_all += batch_shots
                                    completed_batches += 1
                                    if batch_scenes:
                                        last_scene = batch_scenes[-1]
                                        last_shots = last_scene.get('shots', [])
                                        if last_shots:
                                            prev_tail = last_shots[-1].get('end_frame', '') or last_shots[-1].get('subject', '')
                                    yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 完成 ✓ ({batch_shots} 个分镜)"})
                                    # 增量保存：每批完成就落盘
                                    results['shots'] = {
                                        "scenes": all_shot_scenes,
                                        "total_scenes": len(all_shot_scenes),
                                        "total_shots": total_shots_all,
                                        "summary": f"处理中... {completed_batches}/{total_batches} 批",
                                        "directing_notes": ""
                                    }
                                    try:
                                        partial_html = generate_html_report(results, script_name, episode_info=None)
                                        partial_path = _output_path(script_name, '分镜拆解.html', ep_subdir)
                                        with open(partial_path, 'w', encoding='utf-8') as f:
                                            f.write(partial_html)
                                    except Exception:
                                        logging.warning("分镜增量保存失败", exc_info=True)
                                    # 每集独立 HTML：按 shot_id 前缀分组（兼容多种格式）
                                    try:
                                        for ep_pair in batch_eps:
                                            ep_num = ep_pair[0]
                                            if ep_num <= 0:
                                                continue
                                            ep_scenes = []
                                            ep_shot_count = 0
                                            for s in batch_scenes:
                                                ep_shots = [sh for sh in s.get('shots', [])
                                                    if (f'第{ep_num}集' in sh.get('shot_id', '')
                                                        or f'EPISODE {ep_num}' in sh.get('shot_id', '').upper())]
                                                if ep_shots:
                                                    ep_scenes.append({"scene_title": s.get("scene_title", ""), "scene_number": s.get("scene_number", 1), "shots": ep_shots})
                                                    ep_shot_count += len(ep_shots)
                                            if ep_scenes:
                                                ep_result = {"scenes": ep_scenes, "total_scenes": len(ep_scenes), "total_shots": ep_shot_count, "summary": f"第{ep_num}集 · {ep_shot_count} 个分镜", "directing_notes": ""}
                                                ep_html = generate_html_report({"shots": ep_result}, script_name, episode_info=None)
                                                ep_path = _output_path(script_name, f'第{ep_num}集_{time.strftime("%Y%m%d_%H%M%S")}.html')
                                                with open(ep_path, 'w', encoding='utf-8') as f:
                                                    f.write(ep_html)
                                    except Exception:
                                        logging.warning("单集HTML生成失败", exc_info=True)
                                break  # 有结果了就退出重试循环
                            except Exception as e:
                                if attempt == 0:
                                    yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 异常，重试中...（{type(e).__name__}）"})
                                    time.sleep(1)
                                    continue
                                incomplete_batches.append(ep_range)
                                yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"批次 {batch_num}/{total_batches}：{ep_range} 失败 ⚠️ {type(e).__name__}"})
                                if isinstance(e, GeneratorExit):
                                    break
                            break  # 重试也失败了

                    result = {
                        "scenes": all_shot_scenes,
                        "total_scenes": len(all_shot_scenes),
                        "total_shots": total_shots_all,
                        "summary": f"按集分批处理完成（{completed_batches}/{total_batches} 批成功），共 {total_shots_all} 个分镜",
                        "directing_notes": ""
                    }
                    if incomplete_batches:
                        result["incomplete_batches"] = incomplete_batches
                        result["summary"] += f"（⚠️ {len(incomplete_batches)} 批失败：{', '.join(incomplete_batches)}）"
                    if empty_batches:
                        result["empty_batches"] = empty_batches
                        result["summary"] += f"（⚠️ {len(empty_batches)} 批为空：{', '.join(empty_batches)}）"
                    results['shots'] = result
                    # 生成分步 HTML 报告
                    partial_html = generate_html_report(results, script_name, episode_info=None)
                    partial_path = _output_path(script_name, '分镜拆解.html', ep_subdir)
                    with open(partial_path, 'w', encoding='utf-8') as f:
                        f.write(partial_html)
                    yield json_sse("progress", {"step": "shots", "status": "complete", "label": "分镜拆解", "message": f"完成！{completed_batches}/{total_batches} 批，全剧 {total_episodes} 集，生成 {total_shots_all} 个分镜" + (f"（{len(incomplete_batches)} 批失败）" if incomplete_batches else ""), "data": result, "download_url": f"/api/download/{os.path.basename(partial_path)}"})

                else:
                    # 集标记未检测到或只有少量集
                    scene_data = results.get('scenes', {})
                    scene_list = scene_data.get('scenes', []) if isinstance(scene_data, dict) else []
                    sc_count = len(scene_list)
                    ep_info = f"检测到{len(episodes)}集" if episodes else "未检测到集标记"
                    yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"⚠️ {ep_info}，单次处理（共 {sc_count} 个场景）。如需按集拆分，请确保剧本中用 \"第X集\" 或 \"EPISODE X\" 标注每集开始。"})
                    try:
                        shots_api_config = dict(api_config)
                        model_name = shots_api_config.get('model', '')
                        if 'v4-flash' in model_name:
                            shots_api_config['thinking'] = '1'
                        sys_p, user_p = shots.build_list_prompt(script_text, results, style=breakdown_style)
                        raw, _ = _call_ai_retry(sys_p, user_p, shots_api_config, '分镜拆解规划', max_retries=3)
                        shot_plan_raw = raw  # 保存规划步骤的原始输出
                        shot_plan = shots.parse_list(raw)
                        plan_scenes = shot_plan.get('total_scenes', len(shot_plan.get('scenes', [])))
                        if plan_scenes == 0:
                            raw_preview = shot_plan_raw[:500] if shot_plan_raw else "EMPTY"
                            err_msg = "⚠️ 分镜规划返回 0 个场景，无法继续。AI原始返回前500字: " + raw_preview[:300]
                            yield json_sse("progress", {"step": "shots", "status": "error", "label": "分镜拆解", "message": err_msg})
                            results["shots"] = {"scenes":[],"total_scenes":0,"total_shots":0,"summary":"规划步骤失败"}
                            return
                        yield json_sse("progress", {"step": "shots", "status": "processing", "label": "分镜拆解", "message": f"第二轮：生成六模块分镜详情（规划 {plan_scenes} 场）..."})
                        # 情绪时间线预分析
                        et_data = None
                        try:
                            char_names = [c.get('name','') for c in results.get('characters',{}).get('characters',[])]
                            scene_list = results.get('scenes',{}).get('scenes',[])
                            if char_names and scene_list:
                                yield json_sse("progress", {"step": "shots", "status": "processing", "label": "情绪预分析", "message": "分析角色情绪时间线..."})
                                et_sys, et_user = emotion_timeline.build_prompt(script_text, char_names, scene_list)
                                et_raw, _ = _call_ai_retry(et_sys, et_user, api_config, '情绪时间线', max_retries=3)
                                et_data = emotion_timeline.parse_result(et_raw)
                        except Exception as e:
                            logging.warning("情绪时间线分析失败: %s", str(e), exc_info=True)
                        sys_p, user_p = shots.build_detail_prompt(script_text, shot_plan, results, style=breakdown_style, emotion_timeline=et_data)
                        raw, _ = _call_ai_retry(sys_p, user_p, shots_api_config, '分镜拆解详情', max_retries=3)
                        result = shots.parse_result(raw)
                        results['shots'] = result
                        st = result.get('total_shots', 0)
                        if st == 0:
                            plan_info = f"规划步骤共 {plan_scenes} 个场景"
                            detail = '解析结果: total_scenes=' + str(result.get('total_scenes','?')) + ', total_shots=' + str(result.get('total_shots','?')) + ', scenes列表=' + str(len(result.get('scenes',[]))) + '项' 
                            raw_preview = raw[:500] if raw else "EMPTY"
                            yield json_sse("progress", {"step": "shots", "status": "complete", "label": "分镜拆解", "message": f"⚠️ 分镜数为 0。{plan_info}。{detail}。AI原始返回前500字:\n{raw_preview}", "data": result})
                        else:
                            partial_html = generate_html_report(results, script_name, episode_info=None)
                            partial_path = _output_path(script_name, '分镜拆解.html', ep_subdir)
                            with open(partial_path, 'w', encoding='utf-8') as f:
                                f.write(partial_html)
                            yield json_sse("progress", {"step": "shots", "status": "complete", "label": "分镜拆解", "message": f"完成！生成 {st} 个分镜", "data": result, "download_url": f"/api/download/{os.path.basename(partial_path)}"})
                    except Exception as e:
                        yield json_sse("progress", {"step": "shots", "status": "error", "label": "分镜拆解", "message": f"失败：{str(e)}"})
                        if step_filter and step_filter == 'shots':
                            yield json_sse("error", {"message": f"分镜拆解失败: {str(e)}"})
                            return
                        results['shots'] = {"shots": [], "total_shots": 0, "error": str(e)}
            
            # ===== 生成报告（仅自动模式）=====
            if not step_filter:
                yield json_sse("progress", {"step": "report", "status": "processing", "label": "生成报告", "message": "正在生成拆解报告..."})
                
                try:
                    episode_info = detect_episode(script_text)
                    # 确定剧集标签：优先 batch_mode 传入的 episode_label
                    if batch_mode == '1' and episode_label:
                        ep_label = episode_label.replace(' ', '_')
                    elif episode_info:
                        ep_label = f"EP{episode_info['current']}"
                    elif episode_label:
                        ep_label = episode_label.replace(' ', '_')
                    else:
                        ep_label = "EP"
                    
                    # HTML 报告 — batch_mode 时写入剧集子目录
                    html_content = generate_html_report(results, script_name, episode_info)
                    ep_subdir = episode_label if batch_mode == '1' and episode_label else ''
                    html_path = _output_path(script_name, f'{ep_label}_拆解报告.html', ep_subdir)
                    with open(html_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                    
                    # Word 报告
                    word_path = generate_word_report(results, script_name, episode_info)
                    
                    if not os.path.exists(html_path):
                        raise RuntimeError("HTML 报告生成失败")
                    if word_path and not os.path.exists(word_path):
                        raise RuntimeError("Word 报告生成失败")
                    
                    # 资产核对 PPTX — 写入剧集子目录（字段映射为 PPTX 所需格式）
                    pptx_file = None
                    try:
                        assets = _map_extraction_to_pptx_assets(results)
                        ep_num = episode_info['current'] if episode_info else 1
                        pptx_name = f'{script_name}_{ep_label}_资产核对.pptx'
                        pptx_path = _output_path(script_name, pptx_name, ep_subdir) if ep_subdir else os.path.join(os.path.dirname(html_path), pptx_name)
                        build_episode_pptx(assets, script_name, ep_num, pptx_path)
                        pptx_file = os.path.basename(pptx_path)
                    except Exception as e:
                        logging.warning("PPTX 生成失败(非致命): %s", str(e), exc_info=True)
                    
                    # 标记完成 + 批量模式：删除中间分步文件
                    if ep_subdir:
                        _mark_episode_complete(script_name, ep_subdir)
                    if batch_mode == '1' and ep_subdir:
                        _cleanup_intermediate_files(script_name, ep_subdir)
                    
                    yield json_sse("progress", {
                        "step": "report",
                        "status": "complete",
                        "label": "📄 生成报告",
                        "message": f"报告已生成！",
                        "data": {
                            "html_file": os.path.basename(html_path),
                            "word_file": os.path.basename(word_path) if word_path else None,
                        "pptx_file": pptx_file,
                            "script_name": script_name
                        }
                    })
                    
                    # 完成
                    yield json_sse("complete", {
                        "message": "剧本分析全部完成！",
                        "html_file": os.path.basename(html_path),
                        "word_file": os.path.basename(word_path) if word_path else None,
                        "pptx_file": pptx_file,
                        "output_dir": os.path.dirname(html_path) if html_path else config.OUTPUT_DIR,
                        "script_name": script_name,
                        "results": results
                    })
                    
                except Exception as e:
                    yield json_sse("progress", {
                        "step": "report",
                        "status": "error",
                        "label": "📄 生成报告",
                        "message": f"失败：{str(e)}"
                    })
                    yield json_sse("error", {"message": f"报告生成失败: {str(e)}"})
                    return
                
            # 分步模式：生成报告并发送完成事件
            if step_filter:
                try:
                    episode_info = detect_episode(script_text)
                    html_content = generate_html_report(results, script_name, episode_info)
                    if batch_mode == '1' and episode_label:
                        ep_label = episode_label.replace(' ', '_')
                    elif episode_info:
                        ep_label = f"EP{episode_info['current']}"
                    elif episode_label:
                        ep_label = episode_label.replace(' ', '_')
                    else:
                        ep_label = "EP"
                    ep_subdir = episode_label if batch_mode == '1' and episode_label else ''
                    html_path = _output_path(script_name, f'{ep_label}_拆解报告.html', ep_subdir)
                    with open(html_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                    word_path = generate_word_report(results, script_name, episode_info)
                except Exception:
                    logging.warning("最终报告生成失败", exc_info=True)
                    html_path = None; word_path = None
                step_names = {'characters':'角色提取','props':'道具提取','scenes':'场景拆解','shots':'分镜拆解'}
                yield json_sse("complete", {
                    "message": f"{step_names.get(step_filter, step_filter)}完成！",
                    "html_file": os.path.basename(html_path) if html_path else None,
                    "word_file": os.path.basename(word_path) if word_path else None,
                    "output_dir": os.path.dirname(html_path) if html_path else config.OUTPUT_DIR,
                    "script_name": script_name,
                    "results": results
                })
        except Exception as e:
            logging.getLogger("app").error("分析过程出错: %s", str(e))
            yield json_sse("error", {"message": "分析过程出错，请稍后重试"})
    
    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
        }
    )


@app.route('/api/download/<filename>')
def download_report(filename):
    """下载生成的报告"""
    if not _check_rate_limit(_get_client_id(), max_requests=60, window=60):
        return jsonify({"error": "请求过于频繁，请稍后重试"}), 429
    file_path = safe_join(config.OUTPUT_DIR, filename)
    if file_path is None or not os.path.isfile(file_path):
        return jsonify({"error": "文件不存在"}), 404
    
    # 判断文件类型
    if filename.endswith('.docx'):
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        mimetype = 'text/html; charset=utf-8'
    
    return send_file(
        file_path,
        mimetype=mimetype,
        as_attachment=True,
        download_name=os.path.basename(file_path)
    )


@app.route('/api/parse_file', methods=['POST'])
def parse_file():
    """解析上传的文件（.docx/.doc/.pdf/.txt/.md）为纯文本，供前端做多集检测等"""
    if 'file' not in request.files:
        return jsonify({"error": "请上传文件"}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({"error": "文件名为空"}), 400
    
    # 保存临时文件
    ext = os.path.splitext(file.filename)[1].lower()
    tmp_path = os.path.join(config.UPLOAD_FOLDER, f"parse_{uuid.uuid4().hex}{ext}")
    os.makedirs(config.UPLOAD_FOLDER, exist_ok=True)
    file.save(tmp_path)
    
    try:
        text = parse_script(tmp_path)
        # 检测集数
        eps = split_script_by_episodes(text) or []
        episode_markers = []
        for ep_num, ep_text in eps:
            if ep_num > 0:
                episode_markers.append({"number": ep_num, "label": f"第{ep_num}集" if ep_num < 100 else f"EPISODE {ep_num}"})
        
        return jsonify({
            "success": True,
            "text": text,
            "text_length": len(text),
            "episodes": episode_markers,
            "total_episodes": len(episode_markers)
        })
    except Exception as e:
        import traceback
        err_detail = str(e)
        logging.getLogger("app").error("文件解析失败 [%s]: %s\n%s", ext, err_detail, traceback.format_exc())
        return jsonify({"error": f"{ext} 文件解析失败: {err_detail}"}), 500
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass

@app.route('/api/split-shots', methods=['POST'])
def split_shots():
    """将全剧分镜HTML按集拆分为独立文件"""
    data = request.get_json(silent=True) or {}
    script_name = data.get('script_name', '').strip()
    if not script_name:
        return jsonify({"error": "请提供剧本名称"}), 400
    try:
        from tool.shot_splitter import split_shots_by_episode
        result = split_shots_by_episode(script_name)
        return jsonify({"success": True, "folder": result["folder"], "files": result["files"], "count": result["count"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/preview/<filename>')
def preview_report(filename):
    """预览生成的 HTML 报告"""
    file_path = safe_join(config.OUTPUT_DIR, filename)
    if file_path is None or not os.path.isfile(file_path):
        return jsonify({"error": "文件不存在"}), 404
    
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content, 200, {'Content-Type': 'text/html; charset=utf-8'}


@app.route('/api/split_script', methods=['POST'])
def split_script():
    """检测并拆分多集剧本为独立 .docx 文件（不改动原文）"""
    data = request.get_json(silent=True) or {}
    script_text = (data.get('text') or '').strip()
    script_name = (data.get('script_name') or '未命名剧本').strip()
    selected_eps = data.get('episodes')  # [1, 3, 5] 或 None（全选）

    if not script_text:
        return jsonify({"error": "剧本内容为空"}), 400

    script_name = re.sub(r'[\\/:*?"<>|]', '_', script_name)
    script_name = script_name.replace('\x00', '')
    script_name = script_name.lstrip('.-').strip()[:100] or "未命名剧本"

    episodes = split_script_by_episodes(script_text)
    if not episodes:
        return jsonify({"error": "未检测到集标记（第X集 / EPISODE X / EP X），请在剧本中添加集标记后重试"}), 400

    actual_eps = [e for e in episodes if e[0] > 0]
    if not actual_eps:
        return jsonify({"error": "未检测到有效集号"}), 400

    if selected_eps and isinstance(selected_eps, list) and len(selected_eps) > 0:
        selected_set = set(int(s) for s in selected_eps)
        actual_eps = [e for e in actual_eps if e[0] in selected_set]

    if not actual_eps:
        return jsonify({"error": "没有匹配的集号"}), 400

    series = _series_name(script_name)
    output_dir = safe_join(config.OUTPUT_DIR, series)
    if not output_dir:
        return jsonify({"error": "非法剧本名称"}), 400
    os.makedirs(output_dir, exist_ok=True)

    from docx import Document
    from docx.shared import Pt

    results = []
    for ep_num, ep_text in actual_eps:
        ep_label = f"EPISODE {str(ep_num).zfill(2)}"
        ep_dir = os.path.join(output_dir, ep_label)
        os.makedirs(ep_dir, exist_ok=True)

        doc = Document()
        doc.core_properties.author = "foxpaw"
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        doc.add_heading(f'{script_name} — {ep_label}', level=0)
        for para in ep_text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        docx_path = os.path.join(ep_dir, f'{ep_label}_剧本原文.docx')
        doc.save(docx_path)

        results.append({
            "episode": ep_num,
            "label": ep_label,
            "docx_path": docx_path,
            "docx_file": os.path.basename(docx_path),
            "text_length": len(ep_text),
            "text": ep_text
        })

    _init_progress(series, results, selected_eps)
    return jsonify({
        "success": True,
        "series": series,
        "output_dir": output_dir,
        "episodes": results,
        "total": len(results)
    })

@app.route('/api/preprocess', methods=['POST'])
def preprocess():
    """预处理人物小传和故事大纲，生成临时知识库"""
    if not _check_rate_limit(_get_client_id(), max_requests=20, window=60):
        return jsonify({"error": "请求过于频繁，请稍后重试"}), 429
    data = request.get_json(silent=True) or {}
    char_bio = (data.get('char_bio') or '').strip()
    story_outline = (data.get('story_outline') or '').strip()
    script_name = (data.get('script_name') or '未命名剧本').strip()
    api_config = data.get('api_config') or {}
    
    if not char_bio and not story_outline:
        return jsonify({"error": "请提供人物小传或故事大纲"}), 400
    
    script_name = re.sub(r'[\\/:*?"<>|]', '_', script_name)
    script_name = script_name.replace('\x00', '')
    script_name = script_name.lstrip('.-')
    script_name = script_name.strip()[:100] or "未命名剧本"
    
    preprocess_system = """你是剧本分析预处理专家。你的任务是分析人物小传和故事大纲，提取结构化信息。

人物小传分析要求：
- 提取每个角色的名称(name)、别名(aliases)、年龄种族(age_race)、性格标签(personality)、弧光走向(arc)、外貌特征(appearance)、人际关系(relationships)、关键特征(key_traits)

故事大纲分析要求：
- 提取年代/时空设定(time_period)、主要地点(location)、类型标签(genre)、整体基调(tone)、关键剧情道具(key_props)、关键情节点(key_plot_points)、环境风格描述(environment_style)

输出纯JSON："""
    
    preprocess_user = "请分析以下内容并生成结构化知识库。\n\n"
    if char_bio:
        preprocess_user += f"=== 人物小传 ===\n{char_bio}\n\n"
    if story_outline:
        preprocess_user += f"=== 故事大纲 ===\n{story_outline}\n\n"
    
    output_schema = (
        '输出JSON格式：\n'
        '{\n'
        '  "characters": [\n'
        '    {\n'
        '      "name": "角色名",\n'
        '      "aliases": ["别名"],\n'
        '      "age_race": "年龄/种族",\n'
        '      "personality": ["性格标签"],\n'
        '      "arc": "角色弧光走向",\n'
        '      "appearance": "外貌特征",\n'
        '      "relationships": [{"target": "关联角色", "relation": "关系描述"}],\n'
        '      "key_traits": "关键特征"\n'
        '    }\n'
        '  ],\n'
        '  "world": {\n'
        '    "time_period": "年代/时空设定",\n'
        '    "location": "主要地点",\n'
        '    "genre": "类型标签",\n'
        '    "tone": "整体基调",\n'
        '    "key_props": ["关键剧情道具"],\n'
        '    "key_plot_points": ["关键情节点"],\n'
        '    "environment_style": "环境风格描述"\n'
        '  }\n'
        '}\n'
        '只输出JSON。'
    )
    preprocess_user += output_schema
    
    try:
        raw = call_ai(preprocess_system, preprocess_user, api_config)
        parsed = None
        json_text = raw.strip()
        json_start = json_text.find('{')
        json_end = json_text.rfind('}')
        if json_start != -1 and json_end != -1:
            json_text = json_text[json_start:json_end+1]
        try:
            parsed = json.loads(json_text)
        except json.JSONDecodeError:
            return jsonify({"error": "AI 返回格式解析失败: " + raw[:300], "raw": raw[:500]}), 500
        
        series = _series_name(script_name)
        d = safe_join(config.OUTPUT_DIR, series) if series else safe_join(config.OUTPUT_DIR, script_name)
        if not d:
            return jsonify({"error": "非法剧本名称: " + script_name}), 400
        os.makedirs(d, exist_ok=True)
        kb_path = os.path.join(d, f'{series}_人物小传大纲_知识库.json')
        with open(kb_path, 'w', encoding='utf-8') as f:
            json.dump(parsed, f, ensure_ascii=False, indent=2)
        
        char_count = len(parsed.get('characters', []))
        char_names = [c.get('name', '') for c in parsed.get('characters', [])]
        world_info = parsed.get('world', {})
        
        return jsonify({
            "success": True,
            "message": "知识库已生成：" + str(char_count) + " 个角色，" + str(len(world_info)) + " 项世界观信息",
            "summary": {
                "char_count": char_count,
                "char_names": char_names,
                "world": world_info,
            }
        })
    except Exception as e:
        return jsonify({"error": "预处理失败: " + str(e)}), 500


@app.route('/api/get_materials', methods=['POST'])
def get_materials():
    """获取已存储的临时知识库摘要"""
    data = request.get_json(silent=True) or {}
    script_name = (data.get('script_name') or '').strip()
    if not script_name:
        return jsonify({"error": "请提供剧本名称"}), 400
    
    script_name = re.sub(r'[\\/:*?"<>|]', '_', script_name)
    script_name = script_name.replace('\x00', '')
    script_name = script_name.lstrip('.-')
    script_name = script_name.strip()[:100]
    
    series = _series_name(script_name)
    d = safe_join(config.OUTPUT_DIR, series) if series else safe_join(config.OUTPUT_DIR, script_name)
    if not d:
        return jsonify({"has_materials": False}), 200
    
    kb_path = os.path.join(d, f'{series}_人物小传大纲_知识库.json')
    if os.path.exists(kb_path):
        try:
            with open(kb_path, 'r', encoding='utf-8') as f:
                kb = json.load(f)
            char_count = len(kb.get('characters', []))
            char_names = [c.get('name', '') for c in kb.get('characters', [])]
            world_keys = list(kb.get('world', {}).keys())
            return jsonify({
                "has_materials": True,
                "char_count": char_count,
                "char_names": char_names,
                "world_keys": world_keys,
            })
        except Exception:
            return jsonify({"has_materials": False}), 200
    
    return jsonify({"has_materials": False}), 200


@app.route('/api/clear_materials', methods=['POST'])
def clear_materials():
    """清除指定剧本的临时知识库"""
    data = request.get_json(silent=True) or {}
    script_name = (data.get('script_name') or '').strip()
    if not script_name:
        return jsonify({"error": "请提供剧本名称"}), 400
    
    script_name = re.sub(r'[\\/:*?"<>|]', '_', script_name)
    script_name = script_name.replace('\x00', '')
    script_name = script_name.lstrip('.-')
    script_name = script_name.strip()[:100]
    
    series = _series_name(script_name)
    d = safe_join(config.OUTPUT_DIR, series) if series else safe_join(config.OUTPUT_DIR, script_name)
    if d and os.path.exists(os.path.join(d, f'{series}_人物小传大纲_知识库.json')):
        try:
            os.remove(os.path.join(d, f'{series}_人物小传大纲_知识库.json'))
        except OSError:
            pass
    
    return jsonify({"success": True, "message": "临时知识库已清除"})

# ============================================================
# 启动
# ============================================================
@app.route('/api/asset_audit', methods=['POST'])
def asset_audit():
    """资产核对：批量上传分集剧本，提取人物/道具/场景，输出 PPTX + HTML"""
    if not _require_auth():
        return jsonify({"error": "未授权访问"}), 401
    
    uploaded_files = request.files.getlist('files')
    if not uploaded_files:
        return jsonify({"error": "请上传分集文件"}), 400
    
    provider = request.form.get("provider", "ollama")
    api_key = request.form.get("api_key", "")
    model = request.form.get("model", "lfm2:latest")
    
    audit_api_config = {
        "provider": provider,
        "api_key": api_key,
        "model": model,
        "base_url": request.form.get("base_url", "http://localhost:11434"),
        "temperature": "0.3",
        "max_tokens": "4096",
        "timeout": 60
    }
    
    series_name = request.form.get('series_name', '未命名剧集')
    
    file_data = []
    for f in uploaded_files:
        if f.filename:
            ep_match = re.search(r'(\d+)', f.filename)
            ep_num = int(ep_match.group(1)) if ep_match else 999
            content_bytes = f.read()
            ext = os.path.splitext(f.filename)[1].lower()
            file_data.append((ep_num, f.filename, ext, content_bytes))
    
    file_data.sort(key=lambda x: x[0])
    
    def generate():
        try:
            if not file_data:
                yield json_sse("error", {"message": "请上传至少一个分集文件"})
                return
            
            total = len(file_data)
            all_assets = []
            base_dir = os.path.join(config.OUTPUT_DIR, series_name)
            os.makedirs(base_dir, exist_ok=True)
            
            yield json_sse("info", {"message": f"开始资产核对《{series_name}》，共 {total} 集"})
            
            for idx, (ep_num, filename, ext, content_bytes) in enumerate(file_data):
                ep_label = f"EPISODE {str(ep_num).zfill(2)}"
                ep_dir = os.path.join(base_dir, ep_label)
                os.makedirs(ep_dir, exist_ok=True)
                yield json_sse("progress", {
                    "step": "audit", "status": "processing",
                    "label": "资产核对",
                    "message": f"{ep_label}：正在提取资产...",
                    "data": {"episode": ep_num}
                })
                
                tmp_path = os.path.join(config.UPLOAD_FOLDER, f"audit_{idx}{ext}")
                os.makedirs(config.UPLOAD_FOLDER, exist_ok=True)
                with open(tmp_path, 'wb') as wf:
                    wf.write(content_bytes)
                
                try:
                    script_text = parse_script(tmp_path)
                except Exception as e:
                    yield json_sse("progress", {
                        "step": "audit", "status": "error",
                        "message": f"{ep_label}：文件解析失败 - {str(e)}",
                        "data": {"episode": ep_num}
                    })
                    all_assets.append({"characters":[],"props":[],"scenes":[],"error":str(e)})
                    try: os.remove(tmp_path)
                    except OSError: pass
                    continue
                
                try: os.remove(tmp_path)
                except OSError: pass
                
                assets = None
                for retry in range(2):
                    try:
                        assets = extract_assets(script_text, api_config=audit_api_config)
                        break
                    except Exception as ex:
                        if retry == 0:
                            time.sleep(2)
                        else:
                            assets = {"characters":[],"props":[],"scenes":[],"error":str(ex)}
                all_assets.append(assets)
                
                char_count = len(assets.get('characters', []))
                prop_count = len(assets.get('props', []))
                scene_count = len(assets.get('scenes', []))
                
                pptx_path = os.path.join(ep_dir, f'{series_name}_{ep_label}_资产核对.pptx')
                try:
                    build_episode_pptx(assets, series_name, ep_num, pptx_path)
                except Exception as e:
                    yield json_sse("progress", {
                        "step": "audit", "status": "error",
                        "message": f"{ep_label}：PPTX 生成失败 - {str(e)}",
                        "data": {"episode": ep_num}
                    })
                    continue
                
                yield json_sse("progress", {
                    "step": "audit", "status": "complete",
                    "message": f"{ep_label} 完成！人物{char_count} · 道具{prop_count} · 场景{scene_count}",
                    "data": {"episode": ep_num, "characters": char_count, "props": prop_count, "scenes": scene_count}
                })
            
            yield json_sse("progress", {
                "step": "audit", "status": "processing",
                "message": "正在生成全剧资产汇总 HTML..."
            })
            html_path = os.path.join(base_dir, f'{series_name}_全剧资产汇总.html')
            build_summary_html(all_assets, series_name, html_path)
            
            yield json_sse("complete", {
                "message": f"资产核对完成！共 {total} 集",
                "html_file": os.path.basename(html_path),
                "output_dir": base_dir
            })
            
        except Exception as e:
            import traceback
            logging.getLogger("app").error("资产核对失败: %s\n%s", str(e), traceback.format_exc())
            yield json_sse("error", {"message": f"资产核对失败: {str(e)}"})
    
    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'Connection': 'keep-alive'}
    )


@app.route('/api/progress_status', methods=['POST'])
def progress_status():
    """查询全剧分析进度"""
    data = request.get_json(silent=True) or {}
    script_name = (data.get('script_name') or '').strip()
    if not script_name:
        return jsonify({"error": "请提供剧本名称"}), 400
    prog = _load_progress(script_name)
    if not prog:
        return jsonify({"exists": False})
    return jsonify({"exists": True, "progress": prog})


@app.route('/api/progress_delete', methods=['POST'])
def progress_delete():
    """删除进度文件（用户选择从头开始）"""
    data = request.get_json(silent=True) or {}
    script_name = (data.get('script_name') or '').strip()
    if not script_name:
        return jsonify({"error": "请提供剧本名称"}), 400
    try:
        pp = _progress_path(script_name)
        if os.path.isfile(pp):
            os.remove(pp)
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500




if __name__ == '__main__':
    if 'GUNICORN_CMD_ARGS' in os.environ or 'gunicorn' in sys.argv[0]:
        # gunicorn 会自己 import app，不需要 run
        pass
    else:
        # 开发环境：使用 Flask 内置服务器
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        
        print("=" * 60)
        print("  [Film] 剧本拆解大师 v2.52")
        print("=" * 60)
        print("  OpenAI Compatible API: POST /v1/chat/completions")
        print("  Ollama API: POST /api/chat")
        print()
        
        port = int(os.environ.get('PORT', 5000))
        debug = os.environ.get('FLASK_DEBUG', '0') == '1'
        print(f"  开发模式启动: http://localhost:{port}")
        if getattr(sys, 'frozen', False):
            print(f"  Output directory: {os.path.dirname(os.path.abspath(sys.executable))}")
        print("=" * 60)
        
        # 延迟1.5秒后自动打开默认浏览器（PyInstaller兼容）
        def open_browser():
            import time
            time.sleep(1.5)
            subprocess.run(['open', f'http://localhost:{port}'], check=False)
        threading.Thread(target=open_browser, daemon=True).start()
        
        app.run(
            host=config.BIND_HOST,
            port=port,
            debug=debug,
            threaded=True
        )
