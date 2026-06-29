"""
Word 报告生成 - 按知识库输出规则生成 .docx 报告
"""
# (C) foxpaw


import os
import time

MAX_PARA_CHARS = 50000  # 单段落上限，防止超长提示词导致 Word 崩溃


def _safe(text):
    """截断过长文本，防止 Word 文件膨胀/崩溃"""
    s = str(text) if text else ''
    if len(s) > MAX_PARA_CHARS:
        s = s[:MAX_PARA_CHARS] + '\n\n[... 提示词过长已截断，完整内容见 HTML 报告 ...]'
    return s

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config


def generate_word_report(all_results, script_name, episode_info=None):
    """生成 Word 报告（按知识库输出规则）"""
    doc = Document()
    doc.core_properties.author = "foxpaw"
    doc.core_properties.last_modified_by = "foxpaw"
    
    # 版权水印页脚
    _sec = doc.sections[0]
    _ft = _sec.footer
    _ft.paragraphs[0].text = "foxpaw"
    _ft.paragraphs[0].alignment = 0
    for _r in _ft.paragraphs[0].runs:
        _r.font.size = Pt(6)
        _r.font.color.rgb = RGBColor(255, 255, 255)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    
    display_title = script_name
    if episode_info:
        display_title = f"第{episode_info['current']}集 - {display_title}"
    title = doc.add_heading(f'剧本拆解报告 - {display_title}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'生成时间：{time.strftime("%Y-%m-%d %H:%M:%S")}')
    if episode_info:
        doc.add_paragraph(f'内容范围：第{episode_info["current"]}集' + (f' / 共{episode_info["total"]}集' if episode_info.get("total") else ''))
    doc.add_paragraph('—' * 40)
    
    # 1. 角色提取（信息卡 + 提示词）
    if 'characters' in all_results and all_results['characters']:
        result = all_results['characters']
        doc.add_heading('一、角色提取', level=1)
        doc.add_paragraph(f'共识别 {result.get("total_count", 0)} 个角色')
        doc.add_paragraph(f'概述：{result.get("summary", "")}')
        
        for i, char in enumerate(result.get('characters', [])):
            doc.add_heading(f'角色 {i+1}: {char.get("name", "未知角色")}', level=2)
            if char.get('episodes'):
                doc.add_paragraph(f'出场集数：{char["episodes"]}')
            if char.get('age_race'):
                doc.add_paragraph(f'年龄/种族：{char["age_race"]}')
            if char.get('role_type'):
                doc.add_paragraph(f'角色类型：{char["role_type"]}')
            if char.get('description'):
                doc.add_paragraph(f'描述：{char["description"]}')
            if char.get('personality'):
                doc.add_paragraph(f'性格特征：{", ".join(char["personality"])}')
            if char.get('relationships'):
                doc.add_paragraph('关系网络：')
                for rel in char['relationships']:
                    doc.add_paragraph(f'  · {rel.get("target", "")}：{rel.get("relation", "")}')
            if char.get('text_signage'):
                doc.add_paragraph(f'文字/标识：{char["text_signage"]}')
            # 信息卡
            if char.get('info_card'):
                p = doc.add_paragraph()
                run = p.add_run('角色信息卡：')
                run.bold = True
                doc.add_paragraph(str(char['info_card']))
            # 完整提示词
            if char.get('prompt'):
                p = doc.add_paragraph()
                run = p.add_run('中文版提示词：')
                run.bold = True
                doc.add_paragraph(_safe(char['prompt']))
            doc.add_paragraph('')
    
    # 2. 道具提取（信息卡 + 提示词）
    if 'props' in all_results and all_results['props']:
        result = all_results['props']
        doc.add_heading('二、道具提取', level=1)
        doc.add_paragraph(f'共识别 {result.get("total_count", 0)} 个道具（仅统计出现≥2场的道具）')
        doc.add_paragraph(f'概述：{result.get("summary", "")}')
        
        for i, prop in enumerate(result.get('props', [])):
            p = doc.add_paragraph()
            runner = p.add_run(f'道具 {i+1}: {prop.get("name", "")}')
            runner.bold = True
            runner.font.size = Pt(12)
            if prop.get('episodes'):
                doc.add_paragraph(f'  出场集数：{prop["episodes"]}')
            doc.add_paragraph(f'  分类：{prop.get("category", "")} | 出现：{prop.get("frequency", "?")}场 | 数量：{prop.get("quantity", "?")}')
            doc.add_paragraph(f'  描述：{prop.get("description", "")}')
            doc.add_paragraph(f'  材质：{prop.get("material", "")} | 状态：{prop.get("condition", "")} | 年代：{prop.get("period_era", "")}')
            if prop.get('text_signage'):
                p = doc.add_paragraph()
                run = p.add_run(f'  文字/标识：{prop["text_signage"]}')
                run.font.color.rgb = RGBColor(255, 153, 0)
            if prop.get('associated_characters'):
                doc.add_paragraph(f'  关联角色：{", ".join(prop["associated_characters"])}')
            if prop.get('info_card'):
                p = doc.add_paragraph()
                run = p.add_run('  道具信息卡：')
                run.bold = True
                doc.add_paragraph(f'  {prop["info_card"]}')
            if prop.get('prompt'):
                p = doc.add_paragraph()
                run = p.add_run('  中文版提示词：')
                run.bold = True
                doc.add_paragraph(f'  {_safe(prop["prompt"])}')
            doc.add_paragraph('')
    
    # 3. 场景拆解（全景版 + 俯视图版提示词）
    if 'scenes' in all_results and all_results['scenes']:
        result = all_results['scenes']
        doc.add_heading('三、场景拆解', level=1)
        doc.add_paragraph(f'共识别 {result.get("total_count", 0)} 个场景')
        doc.add_paragraph(f'概述：{result.get("summary", "")}')
        
        for scene in result.get('scenes', []):
            doc.add_heading(f'场景 {scene.get("scene_number", "?")}：{scene.get("title", "")}', level=2)
            if scene.get('episode'):
                doc.add_paragraph(f'所属集数：{scene["episode"]}')
            doc.add_paragraph(f'时间：{scene.get("time", "")} | 地点：{scene.get("location", "")} | 类型：{scene.get("scene_type", "")}')
            if scene.get('emotion_tags'):
                doc.add_paragraph(f'情绪标签：{scene["emotion_tags"]} | 光影方案：{scene.get("lighting_scheme", "")}')
            doc.add_paragraph(f'出场角色：{", ".join(scene.get("characters", []))}')
            if scene.get('props'):
                doc.add_paragraph(f'出现道具：{", ".join(scene["props"])}')
            doc.add_paragraph(f'概要：{scene.get("synopsis", "")}')
            doc.add_paragraph(f'戏剧功能：{scene.get("dramatic_function", "")}')
            if scene.get('scene_info_card'):
                p = doc.add_paragraph()
                run = p.add_run('场景信息卡：')
                run.bold = True
                doc.add_paragraph(scene['scene_info_card'])
            # 全景版
            if scene.get('wide_shot_prompt'):
                p = doc.add_paragraph()
                run = p.add_run('全景版提示词（Wide Shot）：')
                run.bold = True
                run.font.color.rgb = RGBColor(255, 153, 0)
                doc.add_paragraph(_safe(scene['wide_shot_prompt']))
            # 俯视图版
            if scene.get('topdown_prompt'):
                p = doc.add_paragraph()
                run = p.add_run('俯视图版提示词（Top-down View）：')
                run.bold = True
                run.font.color.rgb = RGBColor(255, 153, 0)
                doc.add_paragraph(_safe(scene['topdown_prompt']))
            doc.add_paragraph('')
    
    # 4. 分镜拆解（六模块分镜卡）
    if 'shots' in all_results and all_results['shots']:
        result = all_results['shots']
        doc.add_heading('四、分镜拆解', level=1)
        doc.add_paragraph(f'共拆解 {result.get("total_scenes", 0)} 个场景，{result.get("total_shots", 0)} 个分镜')
        doc.add_paragraph(f'概述：{result.get("summary", "")}')
        
        if result.get('directing_notes'):
            doc.add_paragraph(f'导演备注：{result["directing_notes"]}')
        
        for scene_data in result.get('scenes', []):
            doc.add_heading(f'场景：{scene_data.get("scene_title", "")}', level=2)
            
            for shot in scene_data.get('shots', []):
                p = doc.add_paragraph()
                runner = p.add_run(f'分镜 {shot.get("shot_id", shot.get("shot_number", "?"))}')
                runner.bold = True
                runner.font.size = Pt(12)
                doc.add_paragraph(f'  动作链数：{shot.get("action_chains", "?")} | 总时长：{shot.get("total_duration", "?")}s')
                doc.add_paragraph(f'  主体(Subject)：{shot.get("subject", "")}')
                doc.add_paragraph(f'  场景(Setting)：{shot.get("setting", "")}')
                doc.add_paragraph(f'  光影(Light)：{shot.get("lighting", "")}')
                doc.add_paragraph(f'  镜头(Camera)：{shot.get("camera", "")}')
                if shot.get('gaze'):
                    doc.add_paragraph(f'  视线(Gaze)：{shot["gaze"]}')
                if shot.get('prop'):
                    doc.add_paragraph(f'  手持物(Prop)：{shot["prop"]}')
                if shot.get('proximity'):
                    doc.add_paragraph(f'  贴脸距离(Proximity)：{shot["proximity"]}')
                doc.add_paragraph(f'  情绪烈度(Intensity)：{shot.get("intensity", "")}')
                doc.add_paragraph(f'  情绪(Emotion)：{shot.get("emotion", "")}')
                if shot.get('rhythm_notes'):
                    p = doc.add_paragraph()
                    run = p.add_run('  节奏设计：')
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 153, 0)
                    doc.add_paragraph(f'    {shot["rhythm_notes"]}')
                # 秒级动作链
                if shot.get('timeline'):
                    doc.add_paragraph('  秒级动作链：')
                    for t in shot['timeline']:
                        if isinstance(t, dict):
                            doc.add_paragraph(f'    [{t.get("rhythm_role", "")}] {t.get("time_range", "?")} | {t.get("shot_size", "?")}')
                            doc.add_paragraph(f'    动作：{t.get("action", "")}')
                            if t.get('camera_movement'):
                                doc.add_paragraph(f'    镜头运动：{t["camera_movement"]}')
                        else:
                            parts = [str(t[i]) if i < len(t) else '—' for i in range(3)]
                            doc.add_paragraph(f'    {parts[0]} | {parts[1]} | {parts[2]}')
                # 声音与台词
                if shot.get('dialogue'):
                    doc.add_paragraph(f'  台词：{shot["dialogue"]}')
                if shot.get('sfx'):
                    doc.add_paragraph(f'  音效 SFX：{shot["sfx"]}')
                # 收尾画面
                if shot.get('end_frame'):
                    doc.add_paragraph(f'  收尾画面：{shot["end_frame"]}')
                # 统一约束
                if shot.get('constraints'):
                    doc.add_paragraph(f'  统一约束：{_safe(shot["constraints"])}')
                doc.add_paragraph('')
    
    if episode_info:
        output_path = os.path.join(config.OUTPUT_DIR, f'第{episode_info["current"]}集{script_name}.docx')
    else:
        output_path = os.path.join(config.OUTPUT_DIR, f'{script_name}.docx')
    doc.save(output_path)
    return output_path
